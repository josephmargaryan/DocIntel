# src/scraper/file_scrapers.py
import openpyxl
import os
import fitz  # PyMuPDF
import docx
from io import BytesIO
from .base_scraper import BaseScraper
from ..document import Document


class PDFScraper(BaseScraper):
    def scrape(self, filepath: str) -> Document:
        import fitz  # PyMuPDF

        doc = fitz.open(filepath)
        text = ""
        image_paths = []
        # Create a temporary directory specific to this document
        temp_image_dir = os.path.join(
            "temp_images", os.path.splitext(os.path.basename(filepath))[0]
        )
        os.makedirs(temp_image_dir, exist_ok=True)

        for page_num, page in enumerate(doc, start=1):
            text += page.get_text()
            # If page contains images, save them
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image_name = f"image_page{page_num}_{img_index}.png"
                image_path = os.path.join(temp_image_dir, base_image_name)
                pix = fitz.Pixmap(doc, xref)

                try:
                    # Only log important actions
                    if pix.n - pix.alpha < 4:  # this is GRAY or RGB
                        pix.save(image_path)
                    else:  # CMYK or unsupported, convert to RGB
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                        pix.save(image_path)

                    image_paths.append(image_path)

                    # Perform OCR on the image if needed
                    if self.ocr_engine:
                        ocr_text = self.ocr_engine.perform_ocr(image_path)
                        text += ocr_text

                except Exception as e:
                    print(
                        f"[ERROR] Error processing image {img_index} on page {page_num}: {e}"
                    )

        doc.close()
        return Document(text, file_path=filepath, image_paths=image_paths)


class DocxScraper(BaseScraper):
    def scrape(self, filepath: str) -> Document:
        doc = docx.Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return Document("\n".join(full_text))


class ImageScraper(BaseScraper):
    def scrape(self, filepath: str) -> Document:
        if self.ocr_engine:
            text = self.ocr_engine.perform_ocr(filepath)
            image_paths = [filepath]
            return Document(text, file_path=filepath, image_paths=image_paths)
        else:
            raise ValueError("No OCR engine provided for image scraping.")


class ExcelScraper(BaseScraper):
    def scrape(self, filepath: str) -> Document:
        workbook = openpyxl.load_workbook(filepath, data_only=True)
        sheet = workbook.active

        # Initialize text to store extracted data
        text = ""

        # Store data in rows for easy retrieval and CSV export
        rows = []

        # Assume the data starts at row 4 (based on your example)
        for row in sheet.iter_rows(min_row=4, values_only=True):
            if row[0]:  # Assuming row[0] contains "Topic" or relevant data
                text += " ".join([str(cell) for cell in row if cell]) + "\n"
                rows.append(row)

        return Document(text, file_path=filepath, rows=rows)
